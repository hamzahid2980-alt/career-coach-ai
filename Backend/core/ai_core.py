from datetime import datetime
from pathlib import Path
import os
import io
import sys
import json
import re
from typing import Optional, Tuple, List, Dict, Any, Union
from groq import Groq


# Required libraries (ensure they are installed via requirements.txt)
import fitz  # PyMuPDF
import google.generativeai as genai
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from google.api_core import exceptions as google_exceptions


# =========================
# Setup (MODIFIED FOR FALLBACK)
# =========================

API_KEYS = []
MODEL_NAME = "gemini-2.5-flash"

def setup_api_keys():
    """Loads all available Gemini API keys from environment variables."""
    global API_KEYS
    # Try to find .env in current dir or parent dir
    env_path = Path(__file__).resolve().parent.parent / '.env'
    if env_path.exists():
        print(f"DEBUG: Loading .env from {env_path}")
        load_dotenv(dotenv_path=env_path)
    else:
        print(f"DEBUG: .env not found at {env_path}, trying default load_dotenv()")
        load_dotenv()
    
    # Support for comma-separated keys (Render friendly)
    multi_keys = os.getenv("GEMINI_API_KEYS")
    if multi_keys:
        parsed_keys = [k.strip() for k in multi_keys.split(',') if k.strip()]
        API_KEYS.extend(parsed_keys)
        print(f"Loaded {len(parsed_keys)} keys from GEMINI_API_KEYS")

    i = 1
    while True:
        key = os.getenv(f"GEMINI_API_KEY_{i}")
        if i == 1 and not key:
             key = os.getenv("GOOGLE_API_KEY") # Backward compatibility for the first key

        if key:
            if key not in API_KEYS:
                API_KEYS.append(key)
            i += 1
        else:
            break
    
    if not API_KEYS:
        print("CRITICAL ERROR: No 'GEMINI_API_KEY_1' or 'GOOGLE_API_KEY' found in environment variables.")
        # Instead of exit, let's just warn and let the server start (though functions will fail)
        # However, the user might prefer an exit to catch config errors early.
        # I'll keep the exit but make it more descriptive.
        print("Please ensure your .env file exists in the Backend folder and contains GOOGLE_API_KEY.")
        # sys.exit(1) # DISABLED: Don't crash the entire server if AI keys are missing; just let AI fail later.
        pass
    
    print(f"Successfully loaded {len(API_KEYS)} Gemini API key(s).")

# Initialize the keys when the module is loaded
setup_api_keys()

# =========================
# NEW: Central API Call Function with Fallback Logic
# =========================
from core.gemini_handler import GeminiHandler

# Initialize handler
gemini_handler = GeminiHandler()

def _call_gemini_with_fallback(prompt: str, is_chat: bool = False, history: List = None) -> Optional[Any]:
    """
    Calls the Gemini API using the centralized GeminiHandler with fallback mechanism.
    Wrapper to maintain compatibility with existing function calls.
    """
    return gemini_handler.call_gemini(prompt, is_chat, history)

# =========================
# JSON Schema Constants (Your code - UNCHANGED)
# =========================
ASSESSMENT_QUESTIONS_SCHEMA = """
[
  {
    "question_id": "string",
    "question_text": "string",
    "question_type": "single_choice" | "multiple_choice" | "short_answer" | "coding_challenge",
    "options": ["string option 1", "string option 2", "string option 3", "string option 4"],
    "correct_answer_keys": ["string option 1"]
  }
]
"""
ASSESSMENT_EVALUATION_SCHEMA = """
{
  "overall_score": 75,
  "skills_mastered": 3,
  "areas_to_improve": 2,
  "skill_scores": { "Python": 80, "SQL": 60, "Data Analysis": 75 },
  "strengths": ["Demonstrated strong foundational knowledge in Python.", "Understood basic SQL queries."],
  "weaknesses": ["Struggled with complex data manipulation in SQL.", "Limited understanding of advanced data analysis concepts."],
  "recommendations": [
    "Focus on SQL subqueries and window functions for data manipulation.",
    "Practice implementing machine learning algorithms from scratch.",
    "Explore advanced data visualization techniques and tools."
  ]
}
"""
FULL_RESUME_ANALYSIS_SCHEMA = """
{
  "analysis_date": "September 05, 2025",
  "job_role_context": "string",  # CORRECTED: Changed default to 'string' to indicate it's dynamic
  "ai_model": "Google Gemini",
  "overall_resume_score": 68,
  "overall_resume_grade": "Good",
  "ats_optimization_score": 60,
  "professional_profile_analysis": {
    "title": "Professional Profile Analysis",
    "summary": "The candidate presents a clear trajectory of learning and project involvement, demonstrating foundational skills relevant to the target role. However, the profile could benefit from a more concise and impact-driven summary statement tailored directly to a 'Frontend Developer' role, immediately highlighting key value propositions."
  },
  "education_analysis": {
    "title": "Education Analysis",
    "summary": "The education section is concise but could be enhanced. Adding the expected graduation date would be helpful. Listing relevant development coursework (e.g., 'Web Development I & II', 'Data Structures and Algorithms') would reinforce technical expertise. Emphasize any honors or significant academic achievements."
  },
  "experience_analysis": {
    "title": "Experience Analysis",
    "summary": "The projects section is informative but overwhelming. The descriptions are too long and lack quantifiable achievements. Instead of lengthy descriptions, focus on impactful results using numbers and action verbs. For instance, 'Developed a responsive e-commerce platform that increased user engagement by 15%.'"
  },
  "skills_analysis": {
    "title": "Skills Analysis",
    "summary": "The current skills section is adequate but could be structured more effectively for ATS scanning. Consider grouping related skills (e.g., 'Languages: JavaScript, Python', 'Frameworks: React, Angular'). Ensure all frontend-specific skills (e.g., HTML5, CSS3, SASS, Webpack) are explicitly listed and visible."
  },
  "key_strengths": [
    "Diverse Project Portfolio: The candidate has undertaken a wide variety of projects, showcasing initiative and a broad skillset.",
    "Clear Learning Trajectory: Demonstrates continuous learning and application of new technologies.",
    "Foundational Technical Acumen: Possesses a solid understanding of core computer science principles applicable to development."
  ],
  "areas_for_improvement": [
    "Target Role Focus: The resume needs to be sharply focused on the frontend developer role. De-emphasize or remove projects that don't directly showcase relevant frontend skills.",
    "Quantify Achievements: Introduce more quantifiable results (numbers, percentages) in project and experience descriptions.",
    "ATS Keyword Optimization: Integrate common frontend developer keywords (e.g., 'Responsive Design', 'API Integration', 'Version Control', 'TypeScript') more strategically.",
    "Concise Descriptions: Shorten lengthy project/experience descriptions to impactful bullet points.",
    "Consistent Formatting: Ensure consistent formatting across all sections, especially dates and bullet points, for better readability and ATS parsing."
  ],
  "overall_assessment": "This resume demonstrates a strong foundation in computer engineering and relevant project involvement. With targeted refinement to highlight frontend development skills, quantify achievements, and optimize for ATS, the candidate can significantly enhance their chances of securing interviews."
}
"""

# =========================
# Helper Functions (Your code - UNCHANGED)
# =========================
def _safe_json_loads(s: str, fallback=None):
    if not s: return fallback
    s = s.strip()
    if s.startswith("```json"): s = s[7:]
    if s.endswith("```"): s = s[:-3]
    s = s.strip()
    try:
        return json.loads(s)
    except json.JSONDecodeError:
        m = re.search(r"\{.*\}", s, flags=re.DOTALL)
        if m:
            try: return json.loads(m.group(0))
            except json.JSONDecodeError: return fallback
    return fallback

def _norm(s: Optional[str]) -> bool:
    return bool(s and s.strip())

def _smart_join(parts: List[Optional[str]]) -> str:
    return " | ".join([str(p) for p in parts if _norm(p)])

def _best_section_key(target_key: str, available_keys: List[str]) -> Optional[str]:
    if not target_key: return None
    t = target_key.strip().lower().replace(" ", "_").replace("-", "_")
    for k in available_keys:
        k_norm = k.lower().replace(" ", "_")
        if t == k_norm or t in k_norm or k_norm in t: return k
    return None

def parse_user_optimization_input(inp: str) -> Tuple[Optional[str], Optional[str]]:
    val = (inp or "").strip()
    if not val: return None, None
    if ":" in val:
        left, right = val.split(":", 1); return _norm(left), _norm(right)
    if len(val.split()) == 1:
        return val, None
    return None, val

def _stringify_list_content(content: Any) -> str:
    if not isinstance(content, list): return str(content or "")
    string_parts = []
    for item in content:
        if isinstance(item, str): string_parts.append(item)
        elif isinstance(item, dict):
            string_parts.append(", ".join([f"{k.replace('_', ' ').title()}: {v}" for k, v in item.items()]))
        else: string_parts.append(str(item))
    return "\n".join(string_parts)

def extract_text_auto(file_content: bytes, file_extension: str) -> Optional[str]:
    print(f"DEBUG(ai_core): extract_text_auto called for in-memory content (Type: {file_extension})")
    try:
        if file_extension == ".pdf":
            with fitz.open(stream=file_content, filetype="pdf") as doc: 
                return "\n".join([page.get_text() for page in doc])
        elif file_extension == ".docx":
            doc = Document(io.BytesIO(file_content))
            chunks = [p.text for p in doc.paragraphs if _norm(p.text)]
            if doc.tables:
                for table in doc.tables:
                    for row in table.rows:
                        cells_for_chunk = [cell.text for cell in row.cells if _norm(cell.text)]
                        if cells_for_chunk: chunks.append(" | ".join(cells_for_chunk))
            return "\n".join(chunks)
        else:
            return None
    except Exception as e:
        print(f"ERROR(ai_core): Failed to read file content. Exception: {e}", exc_info=True)
        return None

# ============================================
# API Functions (MODIFIED TO USE FALLBACK)
# ============================================

def get_resume_structure(resume_text: str) -> Optional[Dict[str, Any]]:
    prompt = f"""
You are an expert HR Technology engineer specializing in resume data extraction. Your task is to convert the raw text of a resume into a structured, valid JSON object, capturing ALL information with high fidelity.
**Instructions:**
1.  **Use the Base Schema:** For common sections, use the following schema.
2.  **Capture Everything Else:** If you find other sections that do not fit the schema (e.g., "Achievements", "Leadership"), create a new top-level key for them (e.g., "achievements").
3.  **IGNORE THE SKILLS SECTION:** Do not parse the skills section in this step. It will be handled by a different process. Omit the 'skills' key from your output.
**Base Schema:**
{{
  "personal_info": {{ "name": "string", "email": "string", "phone": "string", "linkedin": "string", "github": "string" }},
  "summary": "string",
  "work_experience": [ {{ "role": "string", "company": "string", "duration": "string", "description": ["string", ...] }} ],
  "internships": [ {{ "role": "string", "company": "string", "duration": "string", "description": ["string", ...] }} ],
  "education": [ {{ "institution": "string", "degree": "string", "duration": "string", "description": ["string", ...] }} ],
  "projects": [ {{ "title": "string", "description": ["string", ...] }} ],
  "certifications": [ {{ "name": "string", "description": "string" }} ]
}}
**Critical Rules:**
- If a section from the base schema is NOT in the resume, YOU MUST OMIT ITS KEY from the final JSON. Do not create empty sections.
- Your final output must be a single, valid JSON object starting with `{{` and ending with `}}`. Do not include markdown.
--- RESUME TEXT ---
{resume_text}
--- END RESUME TEXT ---
"""
    response = _call_gemini_with_fallback(prompt)
    if not response: return None
    data = _safe_json_loads(response.text, fallback=None)
    if not data:
        print("\n--- ERROR: GEMINI API FAILED TO RETURN VALID JSON (STRUCTURE) ---")
        return None
    return data

def categorize_skills_from_text(resume_text: str) -> Optional[Dict[str, List[str]]]:
    prompt = f"""
You are an expert technical recruiter and data analyst.
Your sole job is to scan the entire resume text provided and identify only the most relevant, concrete skills.

**Instructions:**
1. All skills explicitly listed in the "Skills" section must be extracted without omission.
2. Additionally, extract other skills if they are explicitly mentioned in experience, projects, education, or certifications.
3. Do not infer or assume skills that are not explicitly stated in the resume.
4. Ignore trivial or non-relevant abilities (e.g., "MS Office", "Internet browsing", unless they are in the Skills section).
5. Categorize the skills into the predefined keys in the JSON schema below.
6. Place each skill in the single most appropriate category.
7. Exclude duplicate or overlapping skills.
8. If a category has no skills, omit the key from the output.

**JSON Output Schema:**
{{
    "Programming Languages": ["Python", "JavaScript", "Java", "C++", ...],
    "Frameworks and Libraries": ["TensorFlow", "PyTorch", "React", "Node.js", "Pandas", ...],
    "Databases": ["MySQL", "PostgreSQL", "MongoDB", ...],
    "Tools and Platforms": ["Git", "Docker", "AWS", "Jira", "Linux", ...],
    "Data Science": ["Machine Learning", "NLP", "Data Visualization", "Predictive Modeling", ...],
    "Soft Skills": ["Leadership", "Teamwork", "Communication", "Problem Solving", ...]
}}

**Critical Rules:**
- Output ONLY the valid JSON object described above.
- Do not add explanations or markdown.

--- RESUME TEXT ---
{resume_text}
--- END RESUME TEXT ---
"""
    response = _call_gemini_with_fallback(prompt)
    if not response: return None
    data = _safe_json_loads(response.text, fallback=None)
    if not data:
        print("\n--- ERROR: GEMINI FAILED TO INFER SKILLS ---")
        return None
    return data


def optimize_resume_json(resume_json: Dict[str, Any], user_input: str, job_description: Optional[str] = None) -> Dict[str, Any]:
    section_req, instruction = parse_user_optimization_input(user_input)
    keys_present = list(resume_json.keys())
    
    job_desc_context = ""
    if job_description and job_description.strip():
        job_desc_context = f"""
        **Job Description Context:**
        Below is the job description for which the resume is being optimized. Incorporate keywords, desired skills, and align the achievements to the requirements of this role.
        ```
        {job_description}
        ```
        """
    base_prompt_context = f"""
CONTEXT: You are an elite career strategist and executive resume writer. Your task is to transform a resume from a passive list of duties into a compelling narrative of achievements that will impress top-tier recruiters.
**Your Transformation Checklist (Apply to every relevant bullet point):**
1.  **Lead with a Powerful Action Verb:** Replace weak verbs with strong, specific verbs (e.g., "Engineered," "Architected," "Spearheaded").
2.  **Quantify Metrics Relentlessly:** Add concrete numbers to show scale and achievement.
3.  **Showcase Impact and Scope:** If a number isn't available, describe the tangible impact or business outcome.
4.  **Integrate Technical Skills Naturally:** Weave technologies into the story of the achievement.
5.  **Ensure Brevity and Clarity:** Remove filler words. Each bullet point should be a single, powerful line.

{job_desc_context} 

**Critical Rules:**
- **Do not modify, add, or delete any titles, names, companies, institutions, or skill names.** This is a strict rule. Only rewrite descriptions.
- DO NOT invent facts or skills.
- DO NOT invent specific numbers.
- Preserve the original data structure.
- Do not modify personal information (name, email, phone).
- Your final output must be only the requested, valid JSON. Do not include markdown.
"""
    if section_req:
        mapped = _best_section_key(section_req, keys_present)
        if not mapped: return resume_json
        sec_data = resume_json.get(mapped)
        prompt = f"""
{base_prompt_context}
TASK: Apply your full transformation checklist to optimize ONLY the following JSON section, named "{mapped}".
--- INPUT JSON SECTION ---
{json.dumps(sec_data, indent=2)}
--- END INPUT JSON ---
"""
    else:
        prompt = f"""
{base_prompt_context}
TASK: Apply your full transformation checklist to optimize all sections of the following resume JSON.
--- FULL INPUT JSON ---
{json.dumps(resume_json, indent=2)}
--- END INPUT JSON ---
"""
    response = _call_gemini_with_fallback(prompt)
    if not response: return resume_json
        
    optimized_data = _safe_json_loads(response.text, fallback=None)
    if not optimized_data:
        print("\n--- ERROR: GEMINI API FAILED TO RETURN VALID JSON (OPTIMIZE) ---")
        return resume_json
            
    if section_req and optimized_data:
        resume_json[mapped] = optimized_data
    elif optimized_data:
        for key, value in optimized_data.items():
            if key in resume_json: resume_json[key] = value

    return resume_json


def optimize_for_linkedin(resume_json: Dict[str, Any], user_input: str, job_description: Optional[str] = None) -> Optional[Dict[str, Any]]:
    context_text = []
    if 'summary' in resume_json: context_text.append(f"Summary:\n{resume_json['summary']}")
    
    all_experiences = resume_json.get('work_experience', []) + resume_json.get('internships', [])
    if all_experiences:
        context_text.append("\nProfessional Experience & Internships:")
        for job in all_experiences:
            description_str = ' '.join(job.get('description', []) if isinstance(job.get('description'), list) else [str(job.get('description', ''))])
            context_text.append(f"- {job.get('role')} at {job.get('company')}: {description_str}")
    
    if 'projects' in resume_json:
        context_text.append("\nProjects:")
        for project in resume_json['projects']:
            description_str = ' '.join(project.get('description', []) if isinstance(project.get('description'), list) else [str(project.get('description', ''))])
            context_text.append(f"- {project.get('title')}: {description_str}")
    
    if 'skills' in resume_json and isinstance(resume_json['skills'], dict):
        skills_summary = ", ".join([f"{cat}: {', '.join(skills)}" for cat, skills in resume_json['skills'].items()])
        context_text.append(f"\nSkills: {skills_summary}")
    
    for key, value in resume_json.items():
        if key not in ['personal_info', 'summary', 'work_experience', 'internships', 'projects', 'skills', 'education', 'certifications', 'resume_metadata', 'raw_text']:
            if isinstance(value, str):
                context_text.append(f"\n{key.replace('_', ' ').title()}:\n{value}")
            elif isinstance(value, list):
                context_text.append(f"\n{key.replace('_', ' ').title()}:\n" + "\n".join([str(item) for item in value]))

    resume_context = "\n".join(context_text)
    section_req, instruction = parse_user_optimization_input(user_input)

    job_desc_context = ""
    if job_description and job_description.strip():
        job_desc_context = f"""
        **Job Description Context:**
        Below is the job description for which the LinkedIn profile is being optimized. Align the content with the keywords, requirements, and tone of this role.
        ```
        {job_description}
        ```
        """
    
    base_prompt_context = f"""
You are an expert LinkedIn profile strategist and personal branding coach.
Your task is to generate compelling, optimized text for a user's LinkedIn profile based on the provided resume content.
**Instructions:**
1.  **Headlines:** Create 2-3 powerful, keyword-rich headline options.
2.  **About (Summary):** Write a compelling, first-person "About" section.
3.  **Experiences:** For EACH job/internship in the context, rewrite the bullet points to be concise and results-oriented.
4.  **Projects:** For EACH project in the context, rewrite its description to be engaging for a LinkedIn audience.

{job_desc_context}
**JSON Output Schema:**
{{
    "headlines": ["string option 1", ...],
    "about_section": "string",
    "optimized_experiences": [ {{ "title": "Role at Company", "description": "string" }} ],
    "optimized_projects": [ {{ "title": "Project Title", "description": "string" }} ]
}}

**Critical Rules:**
- Generate content ONLY from the provided resume context.
- Keep the tone professional but approachable.
- Your final output must be ONLY the valid JSON object that matches the requested task.
"""
    if section_req:
        instr_text = instruction or f"Make the {section_req} section more compelling and professional."
        prompt = f"""
{base_prompt_context}
TASK: Based on the resume context, optimize ONLY the '{section_req}' portion of a LinkedIn profile.
--- RESUME CONTEXT ---
{resume_context}
--- END RESUME CONTEXT ---
"""
    else:
        instr_text = instruction or "Optimize the entire LinkedIn profile, processing every experience and project."
        prompt = f"""
{base_prompt_context}
TASK: Based on the resume context, perform a full optimization of a LinkedIn profile.
--- RESUME CONTEXT ---
{resume_context}
--- END RESUME CONTEXT ---
"""

    response = _call_gemini_with_fallback(prompt)
    if not response: return None
    data = _safe_json_loads(response.text, fallback=None)
    if not data:
        print("\n--- ERROR: GEMINI FAILED TO INFER LINKEDIN CONTENT ---")
        return None
    return data


def _strip_markdown(text: str) -> str:
    """Removes markdown characters like ** and * from a string."""
    # This function is correct and does not interfere with numbered lists.
    return re.sub(r'[\*_`]', '', text)


def get_tutor_explanation(topic: str) -> Optional[Dict[str, Any]]:
    """
    Generates a simple explanation for a technical topic.
    **IMPROVEMENT**: The prompt is now radically simplified to fix the "NA" issue.
    It asks for a single `technical_definition` as a clear paragraph,
    matching the user's example image perfectly.
    """
    prompt = f"""
    Act as a friendly and concise expert tutor for a user who is stuck on the topic: **"{topic}"**.
    Provide a clear explanation in a structured JSON format.

    **JSON OUTPUT INSTRUCTIONS:**
    Generate a JSON object with the following keys. All text must be plain. Do NOT use markdown.

    1.  "analogy": A simple, real-world analogy (1-2 sentences).
    2.  "technical_definition": A single, plain-text paragraph that provides a clear and technically accurate definition. Aim for 3-5 sentences, just like a textbook definition.
    3.  "code_example": A JSON object containing "language" (e.g., "javascript") and "code" (a short, well-commented code snippet). If no code is relevant, the "code" value should be an empty string.
    4.  "prerequisites": An array of 1-3 prerequisite concepts the user might need to know.
    """
    response = _call_gemini_with_fallback(prompt)
    if not response: return None
    cleaned_response_text = response.text.replace('```json', '').replace('```', '').strip()
    try:
        return json.loads(cleaned_response_text)
    except Exception as e:
        print(f"An error occurred in AI Tutor: {e}"); return None



def generate_career_roadmap(user_profile: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    # This function is working correctly from the previous update.
    prompt = f"""
    Act as a world-class AI Career Strategist. Your task is to generate a personalized career action plan as a single, valid JSON object.

    **USER PROFILE:**
    - Current Skills: ```{user_profile.get('current_skills_input')}```
    - Current Proficiency: {user_profile.get('current_level')}
    - Goal: ```{user_profile.get('goal_input')}```
    - Goal Proficiency: {user_profile.get('goal_level')}
    - Time Commitment: {user_profile.get('duration')} duration, {user_profile.get('study_hours')} hours/month.

    **JSON OUTPUT STRUCTURE:**
    Generate a single, valid JSON object. String values within the JSON MUST NOT contain any markdown characters.

    1.  "domain": "Full-Stack Development"
    2.  "extracted_skills_and_projects": {{ "skills": ["React", "Node.js"], "projects": ["E-commerce Platform"] }}
    3.  "job_match_score": {{ "score": 75, "summary": "A brief, plain-text summary, STRICTLY under 50 words." }}
    4.  "skills_to_learn_summary": ["State Management (Redux)", "Advanced API Design"] 
        (STRICT RULE: Focus only on the TOP 6 most critical skills. Never exceed 6 items in this list.)
    5.  "timeline_chart_data": {{ "labels": ["Phase 1: Foundations", "Phase 2: Backend"], "durations": [4, 5] }}
    6.  "detailed_roadmap": [{{ "phase_title": "Phase 1: Frontend Fundamentals", "phase_duration": "4 weeks", "topics": ["React Hooks", "Component Lifecycle"] }}]
        (STRICT RULE: Include exactly 6 high-impact tasks/topics per phase. Each item must be a specific, actionable learning milestone.)
    7.  "suggested_projects": [{{ "project_title": "Real-time Whiteboard", "project_level": "Advanced", "skills_mapped": ["React", "Socket.io"], "what_you_will_learn": "Real-time communication and canvas manipulation.", "implementation_plan": ["Week 1-2: Setup canvas and basic drawing.", "Week 3-4: Implement Socket.io for real-time events."] }}]
    8.  "suggested_courses": [{{ "course_name": "MERN Stack Front To Back", "platform": "Udemy", "url": "https://www.udemy.com/course/mern-stack-front-to-back/", "mapping": "Covers foundational and advanced MERN stack skills for your entire roadmap. This certificate covers the foundational skills in Phase 1 and 2." }}]

    """
    response = _call_gemini_with_fallback(prompt)
    if not response: return None
    cleaned_response_text = response.text.replace('```json', '').replace('```', '').strip()
    try:
        return json.loads(cleaned_response_text)
    except Exception as e:
        print(f"An error occurred during AI roadmap generation: {e}"); return None

def get_chatbot_response(query: str, history: list, career_plan_summary: str) -> dict:
    """
    Generates a chatbot response using the career plan as context.
    **IMPROVEMENT**: The primary rule has been updated to require all responses
    to be formatted as a numbered list (1., 2., 3., ...) instead of bullet points.
    """
    print("AI Core: Received request. The career plan context is a string.")
    cleaned_summary = _strip_markdown(career_plan_summary)

    system_prompt = (
        f"You are an AI career assistant and tutor. Your tone is friendly, concise, and encouraging.\n\n"
        f"**PRIMARY RULE: Always format your entire response as a numbered list (e.g., '1. First point. 2. Second point.').** Do this for every response.\n\n"
        f"**CAREER PLAN CONTEXT:**\n{cleaned_summary}\n\n"
        f"**YOUR TASK:**\n"
        f"First, determine the user's intent from their question.\n\n"
        f"1. **If the user asks ABOUT THEIR CAREER PLAN** (e.g., 'what's in phase 1?', 'what course should I take?'), answer using the Career Plan Context.\n\n"
        f"2. **If the user asks for an EXPLANATION of a technical concept** (e.g., 'what is useEffect?'), provide a short, beginner-friendly explanation.\n\n"
        f"3. **If the question is unrelated**, politely decline with: '1. I can only help with your career plan or related technical topics.'\n\n"
        f"**RESPONSE RULES TO FOLLOW:**\n"
        f"- Format ALL responses as a numbered list (1., 2., 3., ...).\n"
        f"- Do NOT use any other markdown (`**`, `*`, `#`).\n"
        f"- Keep all points brief and easy to understand.\n"
    )

    model_history = []
    for message in history:
        role = 'user' if message.get('role') == 'user' else 'model'
        content = message.get('content', '')
        if content: model_history.append({'role': role, 'parts': [content]})

    full_prompt = f"{system_prompt}\n\nUSER QUESTION: {query}"
    response = _call_gemini_with_fallback(prompt=full_prompt, is_chat=True, history=model_history)

    if not response or not response.text:
        raise Exception("AI response failed after trying all API keys.")

    final_response = _strip_markdown(response.text)
    return {"response": final_response}

def generate_assessment_questions(assessment_type: str, skills: List[str], target_role: Optional[str] = None, num_questions: int = 5, user_id: Optional[str] = None) -> Optional[Dict[str, Any]]:
    """
    Generates a set of assessment questions based on selected skills and target role.
    Uses Gemini 2.5 Flash.
    """
    
    skills_str = ", ".join(skills)
    role_context = f" for a {target_role}" if target_role else ""

    difficulty_hint = "medium difficulty"
    normalized_role = (target_role or "").lower()
    if "junior" in normalized_role:
        difficulty_hint = "beginner to medium difficulty"
    elif "senior" in normalized_role or "lead" in normalized_role:
        difficulty_hint = "medium to advanced difficulty"


    prompt = f"""
    You are an expert technical interviewer and AI assessment designer.
    Your task is to generate a concise, focused skill assessment with exactly {num_questions} questions.
    The assessment should cover the following skills: **{skills_str}**.
    The target context is {assessment_type.replace('_', ' ').title()} role{role_context}, at a {difficulty_hint} level.

    **Instructions for Question Generation:**
    1.  Generate a mix of question types:
        -   **Single-choice (radio buttons):** ~50% of questions. Provide 4 distinct options.
        -   **Multiple-choice (checkboxes):** ~20% of questions. Provide 4 distinct options, clearly indicating ALL correct answers.
        -   **Short-answer:** ~20% of questions. Requires a concise text response.
        -   **Coding challenge:** ~10% of questions. Provide a clear problem statement and expected output/logic. (If this is too complex for 2.5-flash to reliably generate, favor more short-answer).
    2.  Ensure questions cover both theoretical understanding and practical application of the skills.
    3.  Assign a unique `question_id` (e.g., "q1", "q2") to each question.
    4.  For each multiple/single choice question, you MUST provide the `correct_answer_keys` (a list of option values that are correct). This is CRITICAL for automated grading.
    
    **JSON Output Schema (List of Question Objects):**
{ASSESSMENT_QUESTIONS_SCHEMA.strip()}
    **Critical Rules:**
    - Your final output MUST be a JSON array containing exactly {num_questions} question objects.
    - DO NOT include any introductory or concluding text outside the JSON.
    - Ensure `correct_answer_keys` is always a LIST, even if only one answer.
    """

    response = _call_gemini_with_fallback(prompt)
    if not response: return None
    questions = _safe_json_loads(response.text, fallback=None)
    if not questions or not isinstance(questions, list):
        print("\n--- ERROR: GEMINI FAILED TO GENERATE VALID ASSESSMENT QUESTIONS ---")
        return None
    return {"questions": questions}

def evaluate_assessment_answers(user_id: str, submitted_answers: List[Dict[str, Any]]) -> Optional[Dict[str, Any]]:
    """
    Evaluates user's assessment answers using Gemini 2.5 Flash and provides structured results.
    """

    answers_summary = []
    for ans in submitted_answers:
        q_id = ans.get("question_id", "N/A")
        user_response = ans.get("answer")
        
        if isinstance(user_response, list):
            user_response_str = ", ".join(user_response)
        elif user_response is None:
            user_response_str = "No answer provided"
        else:
            user_response_str = str(user_response)
        
        answers_summary.append(f"Question ID: {q_id}\nUser Answer: ```{user_response_str}```\n---")

    answers_text = "\n".join(answers_summary)

    prompt = f"""
    You are an expert technical interviewer and AI grader.
    Your task is to evaluate a user's submitted answers for a skill assessment.
    Provide a comprehensive, structured evaluation based on the answers provided.

    **Instructions for Evaluation:**
    1.  **Calculate Overall Score:** Assign an overall percentage score (0-100%) for the assessment.
    2.  **Identify Skills Mastered/Areas to Improve (Counts):** Based on the questions and answers, estimate how many distinct skills were demonstrated proficiently and how many need significant improvement.
    3.  **List Strengths:** Provide 2-3 specific bullet points highlighting what the user did well.
    4.  **List Weaknesses:** Provide 2-3 specific bullet points highlighting areas where the user struggled or demonstrated gaps.
    5.  **Personalized Recommendations:** Provide 2-3 actionable, general recommendations for improvement. These should be text-based recommendations, not URLs.

    **User's Submitted Answers:**
    {'-'*30}
    {answers_text}
    {'-'*30}

    **JSON Output Schema:**
{ASSESSMENT_EVALUATION_SCHEMA.strip()}
    **Critical Rules:**
    - Your final output MUST be a single, valid JSON object following the schema.
    - DO NOT include any introductory or concluding text outside the JSON.
    - The `skill_scores` should be an object mapping skill names (e.g., Python, SQL) to a proficiency score (0-100). Infer these skills from the context of the assessment.
    """
    response = _call_gemini_with_fallback(prompt)
    if not response: return None
    results = _safe_json_loads(response.text, fallback=None)
    if not results or not isinstance(results, dict):
        print("\n--- ERROR: GEMINI FAILED TO EVALUATE ASSESSMENT ANSWERS ---")
        return None
    return results

def generate_full_resume_analysis(resume_text: str, job_description: Optional[str] = None) -> Optional[Dict[str, Any]]:
    """
    Generates a comprehensive resume analysis report, including overall score,
    ATS score, strengths, areas for improvement, and section-wise feedback.
    """
    job_desc_context = ""
    job_role_hint = "General Candidate"  # Default value

    # --- MODIFIED SECTION ---
    # This block now uses the robust fallback function to get the job role.
    if job_description and job_description.strip():
        job_desc_context = f"""
    The user has provided a job description. Analyze the resume specifically against this job description
    to provide highly tailored feedback, especially for ATS optimization, strengths, and areas for improvement.
    Infer the primary 'Job Role' from this description.

    **Job Description:**
    ```
    {job_description}
    ```
    """
        # A simple AI call to infer job role, now using the fallback mechanism.
        role_prompt = f"Extract the primary job role from the following job description. Respond with only the job role text (e.g., 'Software Engineer', 'Data Scientist', 'Frontend Developer').\n\nJob Description: {job_description}"
        
        role_response = _call_gemini_with_fallback(role_prompt) # Using the new fallback function here.
        
        if role_response and role_response.text:
            inferred_role = role_response.text.strip()
            if inferred_role and len(inferred_role.split()) < 5:  # Basic check for validity
                job_role_hint = inferred_role
        else:
             print(f"Warning: Could not infer job role from JD. Using default.")
    # --- END MODIFIED SECTION ---


    # This is your original main prompt, it remains unchanged.
    prompt = f"""
    You are an expert HR consultant and AI resume analyst. Your task is to provide a comprehensive analysis of the given resume.
    Generate a detailed report covering overall assessment, specific section analyses, key strengths, areas for improvement,
    and a dedicated ATS optimization score, all in a single JSON object.

    **Instructions:**
    1.  **Analysis Date:** Current date (e.g., "September 05, 2025").
    2.  **Job Role Context:** Infer a primary job role from the provided job description (if any) or from the resume itself. Default to "General Candidate" if unclear.
    3.  **AI Model:** "Google Gemini"
    4.  **Overall Resume Score:** A percentage (0-100) reflecting general quality, clarity, and effectiveness.
    5.  **Overall Resume Grade:** A concise word (e.g., "Excellent", "Good", "Fair", "Needs Improvement") corresponding to the score.
    6.  **ATS Optimization Score:** A percentage (0-100) reflecting compatibility with Applicant Tracking Systems, especially considering the job description.
    7.  **Section-wise Analysis:** Provide a 'title' and 'summary' for:
        -   `professional_profile_analysis`: For the summary/objective section.
        -   `education_analysis`: For the education section.
        -   `experience_analysis`: For work experience and projects.
        -   `skills_analysis`: For the skills section.
    8.  **Key Strengths:** 2-3 bullet points highlighting positive aspects.
    9.  **Areas for Improvement:** 3-5 bullet points covering general resume improvements AND specific ATS issues (e.g., keyword gaps, formatting problems).
    10. **Overall Assessment:** A concluding paragraph summarizing the findings and potential for improvement.

    {job_desc_context}

    **Resume Text:**
    ```
    {resume_text}
    ```

    **JSON Output Schema:**
{FULL_RESUME_ANALYSIS_SCHEMA.strip()}
    **Critical Rules:**
    - Your final output MUST be a single, valid JSON object following the schema.
    - DO NOT include any introductory or concluding text outside the JSON.
    - Ensure all scores are integers (0-100).
    - If no job description is provided, make reasonable general assumptions for the 'Job Role Context' and ATS analysis.
    - For `analysis_date`, always use the current date in 'Month DD, YYYY' format.
    - For section summaries, be direct and actionable, similar to the provided examples.
    """
    
    # --- MODIFIED SECTION ---
    # The main API call for the analysis also uses the fallback function now.
    response = _call_gemini_with_fallback(prompt)
    if not response:
        return None # Return None if all API keys fail.
    
    analysis_data = _safe_json_loads(response.text, fallback=None)
    
    if not analysis_data or not isinstance(analysis_data, dict):
        print("\n--- ERROR: GEMINI FAILED TO GENERATE VALID FULL RESUME ANALYSIS ---")
        print("API Response Text:", response.text)
        try: print("API Prompt Feedback:", response.prompt_feedback)
        except ValueError: pass
        print("------------------------------------------------------------------\n")
        return None
    
    # Override job_role_context with the one we inferred earlier.
    if job_role_hint != "General Candidate": 
        analysis_data['job_role_context'] = job_role_hint
    
    # Ensure analysis_date is always current, regardless of what the AI generates.
    analysis_data['analysis_date'] = datetime.now().strftime("%B %d, %Y")

    return analysis_data
    # --- END MODIFIED SECTION ---

def get_interview_chat_response(job_description: str, history: List[Dict[str, str]], difficulty: str) -> Optional[Dict[str, str]]:
    """
    Acts as an AI Interviewer with adjustable difficulty, now with API key fallback.
    """
    # This is your original logic to determine the AI's personality based on difficulty.
    # It remains completely unchanged.
    if difficulty == 'easy':
        personality_prompt = """
        Your Persona: You are a friendly and encouraging hiring manager for an entry-level role.
        Your Goal: Understand the candidate's basic knowledge and potential. Ask foundational, single-topic conceptual questions (e.g., "In Python, what is the difference between a list and a tuple?").
        Your Tone: Supportive and patient.
        Your First Action: Start with a simple, welcoming question like "Thanks for coming in. To start, could you tell me about a project you're proud of that's relevant to this role?"
        """
    elif difficulty == 'hard':
        personality_prompt = """
        Your Persona: You are a sharp, direct senior engineer conducting a final-round interview.
        Your Goal: Rigorously test the candidate's deep technical expertise, problem-solving, and system design skills. Ask challenging, multi-part, or scenario-based questions (e.g., "Given the requirements in the job description, walk me through how you would design a scalable, resilient API for our service. What bottlenecks would you anticipate and how would you mitigate them?").
        Your Tone: Critical, professional, and expecting detailed answers. You will ask tough follow-up questions.
        Your First Action: Start directly with a challenging technical question based on a core skill from the job description.
        """
    else: # Default to Medium
        personality_prompt = """
        Your Persona: You are a professional team lead for a mid-level role.
        Your Goal: Evaluate the candidate's practical competence and fit for the team. Ask a mix of conceptual questions and practical scenarios (e.g., "How would you handle a merge conflict in Git?" or "Explain the concept of 'hoisting' in JavaScript.").
        Your Tone: Professional, direct, and balanced.
        Your First Action: Start with a standard technical screening question.
        """
    # This is your original logic for preparing the chat history.
    # It also remains completely unchanged.
    formatted_history = [{'role': msg['role'], 'parts': [{'text': msg['content']}]} for msg in history]
    system_instruction = f"""
    {personality_prompt}
    
    CRITICAL RULE: You are the INTERVIEWER. The user is the CANDIDATE. You must conduct a realistic interview.
    Base ALL of your questions and analysis strictly on the provided job description context below. Do not ask about skills not mentioned.

    --- JOB DESCRIPTION CONTEXT ---
    {job_description}
    --- END CONTEXT ---
    """
    full_history = [
        {'role': 'user', 'parts': [{'text': system_instruction}]},
        {'role': 'model', 'parts': [{'text': "Understood. I am ready to begin the interview."}]}
    ] + formatted_history
    
    # --- MODIFIED SECTION ---
    # Instead of the try/except block, we now prepare the arguments for our new fallback function.
    
    # The 'prompt' is the newest message from the user.
    last_user_message = full_history[-1]['parts'][0]['text']
    
    # The 'history' is everything that came before the user's newest message.
    chat_history_for_api = full_history[:-1]
    
    # Call our resilient fallback function with the chat parameters.
    response = _call_gemini_with_fallback(
        prompt=last_user_message, 
        is_chat=True, 
        history=chat_history_for_api
    )

    # Check the result and return the appropriate response.
    if not response or not response.text:
        print(f"An error occurred in the interview chat endpoint after all fallbacks.")
        return None # Return None on total failure

    return {"reply": response.text}
    # --- END MODIFIED SECTION ---

def generate_user_comparison(user1_profile: Dict[str, Any], user2_profile: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    """
    Compares two user profiles using Gemini to highlight strengths, skill gaps, and providing a recommendation.
    """
    
    prompt = f"""
    You are an expert HR Talent Analyst. Your task is to compare two candidates based on their profiles and provide a structured comparison.
    
    **Candidate 1:**
    - Name: {user1_profile.get('name', 'N/A')}
    - Email: {user1_profile.get('email', 'N/A')}
    - Total Score: {user1_profile.get('score', 0)}
    - Skills: {user1_profile.get('skills', [])}
    - Stats: {user1_profile.get('stats', {})}
    
    **Candidate 2:**
    - Name: {user2_profile.get('name', 'N/A')}
    - Email: {user2_profile.get('email', 'N/A')}
    - Total Score: {user2_profile.get('score', 0)}
    - Skills: {user2_profile.get('skills', [])}
    - Stats: {user2_profile.get('stats', {})}
    
    **Output JSON Schema:**
    Generate a SINGLE structured JSON object with the following keys. Do NOT use markdown.
    {{
        "common_skills": ["List of skills both have"],
        "user1_distinct_skills": ["Skills unique to Candidate 1"],
        "user2_distinct_skills": ["Skills unique to Candidate 2"],
        "comparison_summary": "A concise paragraph comparing their overall profiles, highlighting who might be better suited for different types of roles.",
        "user1_strengths": ["Key strength 1", "Key strength 2"],
        "user2_strengths": ["Key strength 1", "Key strength 2"],
        "recommendation": "A neutral, objective sentence recommending how to decide between them (e.g., 'Choose Candidate A for frontend-heavy roles, Candidate B for backend focus')."
    }}
    """

    response = _call_gemini_with_fallback(prompt)
    if not response: return None
    
    data = _safe_json_loads(response.text, fallback=None)
    if not data:
        print("\n--- ERROR: GEMINI FAILED TO GENERATE COMPARISON ---")
        return None
        
    return data


def get_interview_summary(job_description: str, history: List[Dict[str, str]], proctoring_data: Optional[Dict] = None) -> Optional[Dict[str, Any]]:
    """
    Analyzes the full interview transcript and provides a performance summary,
    now with special handling for malpractice and forced termination.
    """
    transcript = "\n".join([f"{msg['role']}: {msg['content']}" for msg in history])
    
    # --- NEW: Build a proctoring context for the AI ---
    proctoring_context = ""
    if proctoring_data:
        termination_reason = proctoring_data.get('termination_reason')
        if termination_reason:
            proctoring_context = f"""
            CRITICAL CONTEXT: The interview was terminated early due to malpractice.
            Reason: {termination_reason}
            """
        else:
            # Build a summary of warnings if the interview was not terminated
            warnings = []
            if proctoring_data.get('tab_switch_count', 0) > 0:
                warnings.append(f"{proctoring_data['tab_switch_count']} tab switch(es)")
            if proctoring_data.get('phone_detection_count', 0) > 0:
                warnings.append(f"{proctoring_data['phone_detection_count']} phone detection(s)")
            if proctoring_data.get('no_person_warnings', 0) > 0:
                warnings.append(f"{proctoring_data['no_person_warnings']} instance(s) of no person being detected")
            if proctoring_data.get('multiple_person_warnings', 0) > 0:
                warnings.append(f"{proctoring_data['multiple_person_warnings']} instance(s) of multiple people being detected")
            
            if warnings:
                proctoring_context = f"Proctoring Note: The candidate received warnings for: {', '.join(warnings)}."

    prompt = f"""
    You are an expert career coach and technical recruiter. Your task is to analyze the following mock interview transcript and provide a performance summary.
    
    {proctoring_context}

    **Job Description Context:**
    ```
    {job_description}
    ```

    **Interview Transcript:**
    ```
    {transcript}
    ```

    **Your Analysis Task:**
    Provide a detailed analysis in a valid JSON object with the following keys:
    1.  `"overall_score"`: An integer from 0 to 100.
    2.  `"strengths"`: A list of 2-3 positive points.
    3.  `"areas_for_improvement"`: A list of 2-3 constructive points.
    4.  `"overall_feedback"`: A concise summary paragraph.

    **Critical Rules:**
    - If the "CRITICAL CONTEXT" section indicates the interview was terminated, you MUST:
      1. State the termination reason clearly at the beginning of the `overall_feedback`.
      2. Assign an `overall_score` below 30.
      3. List "Maintaining interview integrity" as the primary area for improvement.
    - If there is a "Proctoring Note", incorporate it into your feedback on professionalism or focus.
    - Your final output must be ONLY the valid JSON object.
    """
    response = _call_gemini_with_fallback(prompt)

    if not response or not response.text:
        print("Error generating interview summary after all fallbacks.")
        return None

    summary_data = _safe_json_loads(response.text, fallback=None)
    
    if not summary_data:
        print("\n--- ERROR: GEMINI FAILED TO GENERATE VALID INTERVIEW SUMMARY ---")
        return None

    return summary_data

def save_resume_json_to_docx(resume_json: Dict[str, Any]) -> Document:
    doc = Document()
    def add_heading(text: Optional[str], level: int = 1):
        t = (text or "").strip(); 
        if t: h = doc.add_heading(t, level=level); h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    def add_para(text: Optional[str], bold: bool = False, style: Optional[str] = None):
        t = (text or "").strip() 
        if t:
            p = doc.add_paragraph(style=style)
            run = p.add_run(t)
            run.bold = bold
            run.font.size = Pt(11)
            if style == "List Bullet": p.paragraph_format.left_indent = Pt(36)
            
    print_order = ['personal_info', 'summary', 'skills', 'work_experience', 'internships', 'projects', 'education', 'certifications']
    
    name_for_title = resume_json.get('personal_info', {}).get('name', '')
    if name_for_title:
        doc.add_heading(name_for_title, level=0)

    contact_info_parts = []
    p_info = resume_json.get('personal_info', {})
    if p_info.get('email'): contact_info_parts.append(p_info['email'])
    if p_info.get('phone'): contact_info_parts.append(p_info['phone'])
    if p_info.get('linkedin'): contact_info_parts.append(p_info['linkedin'])
    if p_info.get('github'): contact_info_parts.append(p_info['github'])
    if contact_info_parts:
        add_para(_smart_join(contact_info_parts))
    
    for section in print_order:
        if section in resume_json:
            content = resume_json[section]
            if section == 'personal_info':
                continue
            
            add_heading(section.replace("_", " ").title(), level=2)
            
            if section == 'summary' and isinstance(content, str):
                add_para(content)
            elif section == 'skills' and isinstance(content, dict):
                for category, skill_list in content.items():
                    if isinstance(skill_list, list) and skill_list:
                        p = doc.add_paragraph();
                        run = p.add_run(category.replace("_", " ").title() + ': '); run.bold = True
                        p.add_run(", ".join(skill_list)); 
            elif isinstance(content, list):
                for item in content:
                    if isinstance(item, str):
                        add_para(item, style="List Bullet")
                    elif isinstance(item, dict):
                        # Ensure any dates/timestamps within item are converted to string before joining/displaying
                        item_copy = item.copy()
                        for k, v in item_copy.items():
                            if isinstance(v, datetime): # CORRECTED: Check for datetime.datetime objects
                                item_copy[k] = v.strftime("%b %d, %Y") # Format as readable string
                        
                        header_parts = [item_copy.get("title"), item_copy.get("name"), item_copy.get("role"), item_copy.get("degree"), item_copy.get("institution")]
                        header = _smart_join(header_parts)
                        if header: add_para(header, bold=True)
                        
                        sub_header_parts = [item_copy.get("company"), item_copy.get("duration")]
                        sub_header = _smart_join(sub_header_parts)
                        if sub_header: add_para(sub_header)
                        
                        desc = item_copy.get("description", [])
                        if isinstance(desc, list):
                            for bullet in desc:
                                if _norm(bullet): add_para(str(bullet), style="List Bullet")
                        elif isinstance(desc, str) and _norm(desc):
                            add_para(str(desc), style="List Bullet")

            elif isinstance(content, str):
                add_para(content)
            
    for section, content in resume_json.items():
        if section not in print_order and section not in ['resume_metadata', 'raw_text','optimized_summary']:
            add_heading(section.replace("_", " ").title(), level=2)
            if isinstance(content, list):
                for item in content: add_para(str(item), style="List Bullet")
            else: add_para(str(content))
            
    print("\n✅ DOCX document generated in memory.")
    return doc

# ai_core.py

# You can REMOVE the 'from google.cloud import speech' import
# You can REMOVE the Firebase Admin initialization and imports if only used for this feature

def get_feedback_on_transcript(transcript: str, question: str, job_description: str) -> Optional[Dict[str, str]]:
    """
    Takes a text transcript and gets feedback and the next question from Gemini.
    This is the new, simpler function that replaces video processing.
    """
    try:
        print(f"DEBUG(ai_core): Getting feedback for transcript: '{transcript}'")
        
        # This prompt is the same as before, but now it's guaranteed to get a transcript
        prompt = f"""
        You are an expert career coach analyzing a mock interview answer.
        Job Description Context: {job_description}
        The question asked was: "{question}"
        The candidate's spoken answer (transcribed) was: "{transcript}"

        Your Task:
        1. Provide brief, constructive feedback on the candidate's answer based on the STAR method, clarity, and relevance.
        2. Generate the next logical interview question.
        JSON Output Schema: {{ "feedback": "string", "next_question": "string" }}
        Rule: Respond ONLY with the valid JSON object.
        """

        gemini_response = _call_gemini_with_fallback(prompt)
        if not gemini_response:
            raise Exception("Gemini call failed during feedback generation.")

        return _safe_json_loads(gemini_response.text)

    except Exception as e:
        print(f"CRITICAL ERROR in get_feedback_on_transcript: {e}")
        return None # Return None on failure
    
def process_audio_answer(audio_content: bytes, question: str, job_description: str) -> Optional[Dict[str, str]]:
    """
    The new, robust pipeline for interview analysis using Groq + Whisper.
    1. Transcribes audio to text using the Whisper-1 model via Groq API.
    2. Sends the high-quality transcript to Gemini for feedback.
    """
    try:
        # Step 1: Transcribe Audio using Whisper-1 via Groq
        print("DEBUG(ai_core): Sending audio content to Groq API for Whisper transcription...")
        groq_client = Groq(api_key=os.getenv("GROQ_API_KEY"))
        
        # We need to wrap the bytes content in a file-like object for the API
        audio_file = ("answer.webm", audio_content, "audio/webm")
        
        transcription = groq_client.audio.transcriptions.create(
            file=audio_file,
            model="whisper-large-v3", # State-of-the-art model
        )
        
        transcript = transcription.text
        if not transcript or not transcript.strip():
            raise ValueError("Transcription result was empty.")
            
        print(f"DEBUG(ai_core): High-quality transcript received: '{transcript}'")

        # Step 2: Send transcript to Gemini for analysis
        prompt = f"""
        You are an expert career coach analyzing a mock interview answer.
        Job Description Context: {job_description}
        The question asked was: "{question}"
        The candidate's spoken answer (transcribed) was: "{transcript}"

        Your Task:
        1. Provide brief, constructive feedback on the candidate's answer based on the STAR method and relevance.
        2. Generate the next logical interview question.
        JSON Output Schema: {{ "feedback": "string", "next_question": "string" }}
        Rule: Respond ONLY with the valid JSON object.
        """

        gemini_response = _call_gemini_with_fallback(prompt)
        if not gemini_response:
            raise Exception("Gemini call failed during feedback generation.")

        return _safe_json_loads(gemini_response.text)

    except Exception as e:
        print(f"CRITICAL ERROR in process_audio_answer: {e}")
        return {
            "feedback": "A technical error occurred while processing your answer. Please try recording again for the same question.",
            "next_question": question
        }
def evaluate_and_adjust_roadmap(current_roadmap: dict, performance_summary: dict, trend_data: Optional[dict] = None) -> dict:
    """
    Analyzes user performance against their current roadmap and adjusts it dynamically.
    Uses collective trend data (this week vs previous weeks) to determine if an upgrade or downgrade is needed.
    """

    trend_context = ""
    if trend_data:
        comp_score = trend_data.get('composite_score', 0)
        trend_context = f"""
    **COLLECTIVE PERFORMANCE & TRENDS:**
    - OVERALL COMPOSITE SCORE: {comp_score}% (Weighted: 30% Progress, 25% Assessments, 25% Interviews, 20% ATS)
    - Assessments: This Week ({trend_data['assessments']['recent']}%) vs Previous ({trend_data['assessments']['prior']}%) | Total Avg: {trend_data['assessments']['total']}%
    - Interviews: This Week ({trend_data['interviews']['recent']}%) vs Previous ({trend_data['interviews']['prior']}%) | Total Avg: {trend_data['interviews']['total']}%
    - ATS Scores: This Week ({trend_data['ats']['recent']}%) vs Previous ({trend_data['ats']['prior']}%) | Total Avg: {trend_data['ats']['total']}%
    - Progress: Tasks Completed This Week ({trend_data['progress']['recent_count']}) vs Previous ({trend_data['progress']['prior_count']}) | Total: {trend_data['progress']['total_count']}
    """

    prompt = f"""
    You are an AI Career Performance Analyst. Your task is to evaluate a user's progress and dynamically adjust their career roadmap.
    
    **CURRENT ROADMAP:**
    `json
    {json.dumps(current_roadmap, indent=2)}
    `
    
    {trend_context}

    **YOUR EVALUATION CRITERIA (Weighted by Composite Score):**
    1. **UPGRADE (Accelerate):** If the *Composite Score* is high (>80%), significantly accelerate the roadmap. Introduce advanced frameworks, complex cloud architecture, or system design projects. Reduce durations of future tasks.
    2. **MAINTAIN:** If the *Composite Score* is between 65-80%, keep the current pace but suggest 1 optional "stretch" task.
    3. **STAGNATION (Extend):** If the *Composite Score* is low due to *Recent Progress* (< 10% on Progress weight) but other scores are high, extend current task durations by 1-2 weeks.
    4. **DOWNGRADE/REFINE (Remedial):** If any *Total Average* is <60% or the *Composite Score* has dropped by >15 pts this week, add foundational "Refresher" tasks to the *current* phase.
    5. **ATS FOCUS:** If the *Latest ATS* score is <70, prioritize 'Resume & LinkedIn Optimization' tasks immediately.

    **OUTPUT:**
    Generate a JSON object which is the UPDATED version of the 'detailed_roadmap' and 'suggested_projects'. 
    You MUST also include a 'performance_feedback' string (max 100 words) summarizing the collective trend and why you made these changes.

    **JSON OUTPUT SCHEMA:**
    {{
        "performance_feedback": "string",
        "is_updated": true,
        "updated_roadmap": {{
             "detailed_roadmap": [...],
             "suggested_projects": [...],
             "skills_to_learn_summary": [...]
        }}
    }}

    **Rules:**
    - Return ONLY the valid JSON object.
    - If no changes are needed, return is_updated: false and the original data.
    - Be encouraging but realistic.
    - **CRITICAL REGENERATION RULE:** When updating, you MUST **completely replace** all UNCOMPLETED tasks with entirely new, logically adaptive topics. Do NOT just reorder or slightly edit existing tasks. They must be fresh learning milestones that reflect the user's current trajectory (Improvement or Struggle).
    - **LIMITS:** 
        - Max 6 items in `skills_to_learn_summary`.
        - Exactly 6 topics per phase in `detailed_roadmap`.
    - **PRESERVATION:** You MUST preserve the names of tasks in 'detailed_roadmap' that have `"is_completed": true` to maintain historical consistency, but you may move them to an earlier "Completed Progress" section or keep them in their original phases if it makes logical sense for the new timeline.

    """
    
    response = _call_gemini_with_fallback(prompt)
    if not response: return None
    
    data = _safe_json_loads(response.text, fallback=None)
    if not data:
        print("\n--- ERROR: GEMINI FAILED TO ADJUST ROADMAP ---")
        return None
        
    return data

def generate_skill_trends_analysis(user_skills: List[str], market_data: List[Dict[str, Any]]) -> Optional[Dict[str, Any]]:
    """
    Analyzes user skills against market trends to provide actionable insights.
    Updated to match Frontend 'trends.js' expectations.
    """
    try:
        # If no user skills, provide generic advice
        skills_text = ', '.join(user_skills) if user_skills else "General Software Engineering"

        prompt = f'''
        You are an expert career strategist. Analyze the following user skills against current market trend data.
        
        **User Skills:** {skills_text}
        
        **Market Trend Data:**
        {json.dumps(market_data, indent=2)}
        
        **Task:**
        1. create a concise "Analysis Summary" (2 sentences) about how the user's skills align with the market.
        2. Identify 3 specific "Recommendations" for skills to learn or improve.
        
        **Output JSON Format (Strict):**
        {{
            "analysis_summary": "Your skills in X are strong, but the market is shifting towards Y...",
            "recommendations": [
                {{
                    "skill": "Skill Name", 
                    "trend_relevance": "Why this is hot (e.g., 'High demand in AI roles')", 
                    "learning_path": "How to start (e.g., 'Build a project using framework Z')"
                }}
            ]
        }}
        '''
        
        response = _call_gemini_with_fallback(prompt)
        if not response: return None
        
        return _safe_json_loads(response.text, fallback=None)
    except Exception as e:
        print(f"Error in generate_skill_trends_analysis: {e}")
        return None

def extract_event_details(subject: str, body: str, date: Any = None) -> Dict[str, Any]:
    """
    Extracts event details (Interview/Hackathon) from email content using Gemini.
    Returns a dictionary suitable for Google Calendar/Tasks.
    """
    prompt = f"""
    Analyze the following email to extract event details for a developer's schedule.
    Identify if it's an 'Interview' or a 'Hackathon'. If neither, return null.
    
    Email Subject: {subject}
    Email Body: {body}
    Reference Date: {date if date else 'Today'}

    Output strictly in JSON format:
    {{
        "title": "Short event title (e.g., 'Interview with Google', 'HackMIT 2025')",
        "type": "Interview" or "Hackathon",
        "start_time": "ISO 8601 datetime string (YYYY-MM-DDTHH:MM:SS) or null if not found",
        "end_time": "ISO 8601 datetime string or null (assume 1 hour for interviews if not specified)",
        "description": "Brief summary of the event",
        "preparation_tasks": ["List of 2-3 key topics to prep for this specific event"]
    }}
    """
    try:
        response = _call_gemini_with_fallback(prompt)
        text_content = response.text if response and hasattr(response, 'text') else str(response)
        # Clean specific markdown wrapping if present
        clean_text = text_content.replace("```json", "").replace("```", "").strip()
        return json.loads(clean_text)
    except Exception as e:
        print(f"Error extracting event details: {e}")
        return {}

def draft_application_email(job_desc: str, resume_summary: str, email_type: str, user_name: str = None) -> str:
    """
    Drafts a job application, cold email, or follow-up email.
    """
    sign_off_name = user_name if user_name else "[Your Name]"
    
    prompt = f"""
    You are an expert career coach acting as a specialized email copywriter.
    Draft a professional, compelling, and concise '{email_type}' email.
    
    Context:
    - Target Job Description/Company Info: "{job_desc}"
    - Candidate's Key Qualifications (Resume Summary): "{resume_summary}"
    - Sender Name: "{sign_off_name}"

    Instructions:
    - Tone: Professional, confident, but not arrogant.
    - Structure: Clear subject line (if applicable), personalized opening, value proposition connecting candidate skills to job needs, and a clear call to action.
    - For 'Cold Email': Focus on value add and brevity.
    - For 'Follow-up': Be polite, reiterate interest, and add a small new value point if possible.
    - For 'Application': Standard cover letter style but optimized for email body.
    - Output ONLY the email content. If a subject line is needed, put it on the first line prefixed with 'Subject: '.
    """
    
    try:
        response = _call_gemini_with_fallback(prompt)
        if response and hasattr(response, 'text'):
            return response.text
        elif response and isinstance(response, str):
            return response
        return "Error: Empty response from AI."
    except Exception as e:
        print(f"Error drafting email: {e}")
        return f"Subject: Application Request\n\nDear Hiring Team,\n\nPlease accept this email as my application. (Error generating full draft: {str(e)})\n\nSincerely,\n{sign_off_name}"

def analyze_interview_feedback(current_analysis: Dict, feedback: str) -> Dict[str, Any]:
    """
    Analyzes raw interview notes/feedback to update the user's cumulative interview profile.
    Generates radar chart scores and lists strengths/weaknesses.
    """
    prompt = f"""
    Analyze the following raw notes/feedback from a recent technical interview.
    Update the cumulative interview profile of the candidate.

    Raw Feedback/Notes:
    "{feedback}"

    Current Cumulative Stats (for context, do not just repeat these, evolve them based on new data):
    {json.dumps(current_analysis.get('skill_scores', {}), indent=2)}

    Output strictly in JSON format:
    {{
        "cumulative_advice": "A short, synthesized advice paragraph (2-3 sentences) focusing on the most critical trend across interviews.",
        "skill_scores": {{
            "Technical Knowledge": <int 0-100>,
            "Communication": <int 0-100>,
            "Confidence": <int 0-100>,
            "Problem Solving": <int 0-100>,
            "System Design": <int 0-100>
        }},
        "latest_interview_topics": ["List of technical topics covered in this specific session"],
        "weaknesses": ["List of top 3 weaknesses identified in this session"],
        "strengths": ["List of top 3 strengths identified in this specific session"]
    }}
    """
    try:
        response = _call_gemini_with_fallback(prompt)
        text_content = response.text if response and hasattr(response, 'text') else str(response)
        clean_text = text_content.replace("```json", "").replace("```", "").strip()
        return json.loads(clean_text)
    except Exception as e:
        print(f"Error analyzing interview feedback: {e}")
        return {
            "cumulative_advice": "Could not analyze recent feedback due to an error.",
            "skill_scores": current_analysis.get('skill_scores', {
                "Technical Knowledge": 50, "Communication": 50, "Confidence": 50, "Problem Solving": 50, "System Design": 50
            }),
            "latest_interview_topics": [],
            "weaknesses": ["Error parsing feedback"],
            "strengths": []
        }
